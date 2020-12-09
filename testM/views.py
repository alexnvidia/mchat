from django.shortcuts import render, get_object_or_404, redirect, HttpResponseRedirect
from django.utils import timezone
from django.contrib.auth.decorators import login_required
from django.contrib.sites.shortcuts import get_current_site
from django.utils.encoding import force_text
from django.contrib.auth import login, authenticate
from .models import Mchat , Item, Patient, FollowUpItem, Profile,Patient_historic
from .forms import mchatForm,mchatTest,PatientForm,mchatFollowup,profileForm,custom_question,adicional_info
from .forms import SignUpForm
from django.forms import formset_factory, modelformset_factory
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.contrib.auth.models import User
from django.views.generic.list import ListView
from django.views.generic.edit import CreateView, UpdateView, DeleteView, FormMixin
from django.views.generic.base import TemplateView
from django.views.generic.detail import DetailView
from django.urls import reverse, reverse_lazy
from django.contrib.admin.views.decorators import staff_member_required
from django.utils.decorators import method_decorator
from django.core import serializers
from django.core.files.storage import FileSystemStorage
from django.http import HttpResponse,JsonResponse
from django.template.loader import render_to_string
from django.template import RequestContext
from django.db import IntegrityError
from django.utils.http import urlsafe_base64_decode
from django.utils.encoding import force_bytes
from django.utils.http import urlsafe_base64_encode
from .tokens import account_activation_token
from django.core.mail import send_mail
from django.core.mail import EmailMultiAlternatives
from django.utils.html import strip_tags
from django.conf import settings
from weasyprint import HTML, CSS
from datetime import date
from zipfile import ZipFile
from io import BytesIO, StringIO
from docx import Document
from docx.shared import Inches
import io
import time




# Pdb Import
# Create your views here.

"""
class StaffRequiredMixin(object):
	
	Mixin que solo permite acceso a los usuarios que son
	de tipo staff

	def dispatch(self, request, *args, **kwargs):
		if not request.user.is_staff:
			return redirect(rever)
"""
""" Costants"""

messageFinalScoreNoFollow = "No es necesario hacer la entrevista de seguimiento o segunda etapa, los resultados no indican riesgo de TEA."
messageFinalScoreFollow = "Es recomendable hacer la entrevista de Seguimiento o segunda etapa para poder tener un resultado completo."
messageFinalScoreNoFollowRisk = "Puede prescindir de la entrevista de seguimiento o segunda etapa."
TRIGGER = "TRIGGER"
GROUP = "GROUP"
SINGLE = "SINGLE"
PASA = "PASA"
NOPASA = "NOPASA"
GROUP = "GROUP"
AUDIT = "AUDIT"


""" Constants """


class MchatListView(ListView):

    model = Mchat        

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['now'] = timezone.now()
        return context

@method_decorator(login_required, name='dispatch')
class PatientListView(ListView):

	model = Patient

	def get_context_data(self, **kwargs):
		context = super().get_context_data(**kwargs)
		context['now'] = timezone.now()
		return context
	def get_queryset(self):
		supervisor_filter=User.objects.get(username=self.request.user)
		return Patient.objects.filter(supervisor=supervisor_filter)


@method_decorator(login_required, name='dispatch')
class PatientHistoricView(ListView):

	model = Patient_historic

	def get_context_data(self, **kwargs):
		context = super().get_context_data(**kwargs)
		context['now'] = timezone.now()
		return context
	def get_queryset(self):
		return Patient_historic.objects.filter(patient = self.kwargs['pk'])



""" para que el PatientForm sea válido debo asignar un valor a la
clave foranea supervisor ya que esta no puede ser nula
y se asume que el usuario logueado es el supervisor de ese 
paciente"""
@method_decorator(login_required, name='dispatch')
class PatientCreate(CreateView):
	model = Patient
	form_class = PatientForm
	template_name_suffix = '_create_form'
	success_url = reverse_lazy('mchats:patients')

	def form_valid(self, form):
		user = User.objects.get(username=self.request.user)
		form.instance.supervisor = user
		return super(PatientCreate, self).form_valid(form)


	

@method_decorator(staff_member_required, name='dispatch')
class MchatUpdate(UpdateView):
    model = Mchat
    form_class = mchatForm
    template_name_suffix = '_update_form'

    def get_success_url(self):
        return reverse_lazy('mchats:update', args=[self.object.id])

@method_decorator(login_required, name='dispatch')
class ProfileUpdate(UpdateView):
	model = Profile
	form_class = profileForm
	template_name_suffix = '_update_form'	


	def get_success_url(self):
		return reverse_lazy('mchats:mchats')

	def get_object(self):
		return Profile.objects.get(user=self.kwargs['pk'])



@method_decorator(login_required, name='dispatch')
class PatientUpdate(UpdateView):
	model = Patient
	form_class = PatientForm
	template_name_suffix = '_update_form'

	def form_valid(self, form):
		user = User.objects.get(username=self.request.user)
		form.instance.supervisor = user
		return super(PatientUpdate, self).form_valid(form)


	def get_success_url(self):
		return reverse_lazy('mchats:patients')

@method_decorator(login_required, name='dispatch')
class PatientDelete(DeleteView):
	model = Patient
	success_url = reverse_lazy('mchats:patients')


@method_decorator(login_required, name='dispatch')
class ProfileDetailView(DetailView):
	model = Profile

	def get_context_data(self, **kwargs):
		context = super().get_context_data(**kwargs)
		context['now'] = timezone.now()
		return context

	def get_object(self):
		return Profile.objects.get(user=self.kwargs['pk'])

	def get_queryset(self):
		return Profile.objects.filter(user=self.kwargs['pk'])





def MchatScore(question_id,option):
		puntuacion = 0
		if(question_id == 2 or question_id == 5 or question_id == 12 ):
			if(option == "True"):
				puntuacion = 1
		else:
			if(option == "False"):
				puntuacion = 1
		return puntuacion

def FinalMchatRfScore(puntuacion_mchat):
	if(puntuacion_mchat >= 2):
		return True
	return False


def transform_date_to_age(days):
	age = ""
	if (days > 365):
		year = days//365
		if(year > 1):
			year_str = str(year) + " años"
			age = year_str
		else:
			year_str = str(year) + " año"
			age = year_str
		rest = days%365
		if (rest >= 30):
			month = rest//30
			if (month > 1):
				month_str = " y " + str(month) + " meses"
				age+=month_str
			else:
				month_str = " y " + str(month) + " mes"
				age+=month_str
		else:
			days_rest = rest
			if (days_rest) > 1:
				age+= " y " + str(days_rest) + " días"

		return age
	else:
		month = days//30
		if (month > 1):
			month_str =str(month) + " meses"
			age = month_str
		else:
			month_str = str(month) + " mes"
			age = month_str

		days_rest = days%30
		if (days_rest > 1):
			age += " y " + str(days_rest) + " días"
		else:
			age += " y " + str(days_rest) + " día"
		return age
	return age





def str_to_bool(string):
	if(string == "True"):
		return True
	else:
		return False


def parse_custom_quest(question,final_quest,question_id):
	final_str = ""
	list_quest = question.split(', ')
	list_aux = []

	#de la lista extraigo los item de las cuestiones para luego ordenarlas
	if question != '':
		for l in list_quest:
			if l[0:2].find('.') == -1:
				list_aux.append(int(l[0:2]))
			else:
				list_aux.append(int(l[0:1]))
		print(type(question_id))
		list_aux.append(question_id)
		list_aux.sort()
		list_quest.insert(list_aux.index(question_id),str(question_id) + '.' + final_quest)
		if len(list_quest) == 1:
			for q in list_quest:
				final_str = q
			return final_str
		elif len(list_quest) > 1:
			for q in list_quest:
				if list_quest.index(q) == len(list_quest) - 1:
					final_str+= q
				else:
					final_str+= q + ', '
			return final_str
	else:
		final_str = str(question_id) + '.' + final_quest
		return final_str
	return final_str

def split_validate(question,final_quest,question_id):
	list_quest = question.split(', ')
	final_quest = str(question_id) + '.' +  final_quest
	bool_message = True
	for l in list_quest:
		if(l[0:len(str(question_id))] == final_quest[0:len(str(question_id))]):
			bool_message = False
			return bool_message
	return bool_message

def split_validate_undo(question,question_id):
	bool_message = False

	if len(question) > 0:
		for c in question:
			if(c[0:len(str(question_id))] == str(question_id)):
				bool_message = True
	return bool_message


def undo_custom(question,question_id):
	final_str = ""
	element = ""
	for q in question:
		if q[0:len(str(question_id))] == str(question_id):
			element = q

	if element in question:
		question.pop(question.index(element))
	else:
		return final_str
	if len(question) == 1:
		for l in question:
			final_str = l 
	elif len(question) > 1:
		for l in question:
			if question.index(l) == len(question) - 1:
				final_str+= l
			else:
				final_str+= l + ', '
	return final_str




def option_to_score_list(option,lista_score):
	if (option == "True"):
		lista_score+='1'
		return lista_score
	elif (option == "False"):
		lista_score+='0'
		return lista_score
	else:
		lista_score+='X'
		return lista_score
	return lista_score

def option_to_n(option):
	n = 0
	if(option == "True"):
		n = 1
		return n
	else:
		n = 0
		return n
	return n

def construct_followup_list(question_id,followup_list):
	
	if (question_id > 0 and question_id <=9):
		followup_list+=('0'+str(question_id))
		return followup_list
	else:
		followup_list+=str(question_id)
		return followup_list
	return followup_list


def add_all_optionRF(total_page,request):
	item_scoreRF = ""
	for n in range(1,total_page+1):
		item_scoreRF+=request.session[str(n)]
		request.session[str(n)] =""
	return item_scoreRF




def char_to_list(followup_list):
	listf=[]
	charBak=""

	for c in followup_list:
		charBak+=c
		if(len(charBak) == 2):
			listf.append(int(charBak))
			charBak=""
	return listf


def item_risk2list(item_risk,Item_list):
	listr = []
	j = 0
	if len(item_risk) == len(Item_list):
		for i in Item_list:
			if(item_risk[j] == '1'):
				listr.append(i)
				j+=1
			else:
				j+=1
	return listr

def select_c_quest(custom_quest, question_id):
	element = "noresult"
	for c in custom_quest:
		if (len(c) > 0 and c[0] == str(question_id)):
			element = c
	return element






def objetc_to_list(listaFollowUp):
	followup_object=[]

	for l in listaFollowUp:
		followup_object.append(Item.objects.get(question_id=l))
	return followup_object






def IfMchatFollowUp(puntuacion_mchat):
	resultDict = {
	  "message": "none",
	  "result": False,
	  "finish": False
	  
	}
	if (puntuacion_mchat >= 0 and puntuacion_mchat <=2):
		resultDict["result"] = False
		resultDict["message"] = messageFinalScoreNoFollow
		resultDict["finish"] = True
		return resultDict
	elif (puntuacion_mchat >= 3 and puntuacion_mchat <=7):
		resultDict["result"] = True
		resultDict["message"] = messageFinalScoreFollow
		resultDict["finish"] = False
		return resultDict
	elif (puntuacion_mchat >= 8 and puntuacion_mchat <=20):
		resultDict["result"] = True
		resultDict["message"] = messageFinalScoreNoFollowRisk
		resultDict["finish"] = True
		return resultDict
	return resultDict 

def MchatScoreRf(question_id,count_pasa,list_trigger,count_nopasa,count_group,list_single):
	score = 1 # 0 es PASA y 1 es NOPASA
	if (question_id == 1 or question_id == 10 or question_id == 11 or question_id == 16):
		if(count_pasa > count_nopasa):
			score = 0
			return score
		elif (count_pasa < count_nopasa):
			score = 1
			return score
		else:
			return score
	elif (question_id == 2):
		if(count_group == 0):
			score = 0
			return score
		else:
			score = 1
			return score
	elif (question_id == 3 or question_id == 4):
		if (count_pasa > 0):
			score = 0
			return score
		else:
			score = 1
			return score
	elif (question_id == 5):
		if (count_pasa > 0 or count_nopasa == 0):
			score = 0
			return score
		elif (count_nopasa > 0 and list_trigger[0] == 0):
			score = 0
			return score
		else:
			score = 1
			return score
	elif (question_id == 6):
		if (count_group > 0 and list_trigger[0] == 1):
			score = 0
			return score
		else:
			score = 1
			return score
	elif (question_id == 7):
		if (count_pasa > 0 and list_trigger[0] == 1 and list_trigger[1] == 1):
			score = 0
			return score
		else:
			score = 1
			return score
	elif (question_id == 8):
		if (list_single[0] == 1 and count_pasa > 0 and list_trigger[0] == 1):
			score = 0
			return score
		else:
			score = 1
			return score
	elif (question_id == 9):
		if (count_pasa > 0 and list_trigger[0] == 1):
			score = 0
			return score
		else:
			score = 1
			return score
	elif (question_id == 12):
		if (count_group >= 2 and count_pasa > count_nopasa):
			score = 0
			return score
		elif (count_group < 2):
			score = 0
			return score
		else:
			score = 1
			return score
	elif (question_id == 13):
		if (list_single[0] == 1):
			score = 0
			return score
		else:
			score = 1
			return score
	elif (question_id == 14):
		if (count_pasa >= 2):
			score = 0
			return score
		elif (count_pasa == 1 and list_trigger[0] == 1 and list_trigger[1] == 1):
			score = 0
			return score
		else:
			score = 1
			return score
	elif (question_id == 15):
		if (count_group >= 2):
			score = 0
			return score
		else:
			score = 1
			return score
	elif (question_id == 17 or question_id == 20):
		if (count_group > 0 ):
			score = 0
			return score
		else:
			score = 1
			return score
	elif (question_id == 18):
		if (list_single[0] == 0 and list_trigger[0] == 1 and count_group > 0):
			score = 0
			return score
		elif (list_single[0] == 1 and count_group > 0):
			score = 0
			return score
		else:
			score = 1
			return score
	elif (question_id == 19):
		if (list_single[0] == 1):
			score = 0
			return score
		elif (list_single[0] == 0 and list_trigger[0] == 0 and list_trigger[1] == 0):
			score = 1
			return score
		else:
			score = 0
			return score
	return score

def list_custom_n(list_question):
	list_auxn = []
	for l in list_question:
			if l[0:2].find('.') == -1:
				list_auxn.append(int(l[0:2]))
			else:
				list_auxn.append(int(l[0:1]))
	return list_auxn




def set_option_to_rf(queryset,question,item_scoreRF):
	i = 0
	j = 0
	y = 0
	z = 0
	lista_rf_def = []
	lista_option = []
	lista_quest = question.split(', ')
	list_aux = []
	list_auxn = []
	list_questn = []
	list_position = []

	if(len(item_scoreRF) > 0 and len(item_scoreRF) == queryset.count()):
		for q in queryset:
			if (item_scoreRF[i] == '1'):
				lista_rf_def.append(q.pk)
				lista_option.append(True)
				i+=1
			elif (item_scoreRF[i] == '0'):
				lista_rf_def.append(q.pk)
				lista_option.append(False)
				i+=1
			elif (item_scoreRF[i] == '2' or item_scoreRF[i] == '3' or item_scoreRF[i] == '4'):
				lista_rf_def.append(q.pk)
				lista_option.append(item_scoreRF[i])
				i+=1
			else:
				i+=1
		queryset = queryset.filter(pk__in=lista_rf_def)

		for c in queryset:
			if c.question.find("Otro(describa)") != -1:
				list_aux.append(str(c.item) + '. '+ "Otro(describa)")
		if question != '':
			list_auxn = list_custom_n(list_aux)
			list_questn = list_custom_n(lista_quest)
		print(list_auxn)
		print(list_questn)
		if len(list_auxn) != len(list_questn) and len(list_questn) <= len(list_auxn):
			for r in list_questn:
				if r in list_auxn:
					list_position.append(list_auxn.index(r))
			for c in lista_quest:
				list_aux[list_position[z]] = c
				z+=1
			lista_quest = list_aux
		else:
			lista_quest = lista_quest





		for q in queryset:		
			q.option = lista_option[j]
			j+=1
			if question != '':
				if q.question.find("Otro(describa)") != -1:
					if lista_quest[y][0:2].find('.') == -1 and str(q.item) == lista_quest[y][0:2]:
						q.question=q.question.replace(q.question[q.question.find('.')+2:len(q.question)],lista_quest[y][3:len(lista_quest[y])])
					elif lista_quest[y][0:2].find('.') != -1 and str(q.item) == lista_quest[y][0:1]:
						q.question=q.question.replace(q.question[q.question.find('.')+2:len(q.question)],lista_quest[y][2:len(lista_quest[y])])
					y+=1


	return queryset

def set_option_to_rf_nf(queryset,item_scoreRF):
	i = 0
	j = 0
	lista_option = []

	if(len(item_scoreRF) > 0 and len(item_scoreRF) == queryset.count()):
		for q in queryset:
			if (item_scoreRF[i] == '1'):
				lista_option.append(True)
				i+=1
			elif (item_scoreRF[i] == '0'):
				lista_option.append(False)
				i+=1
			elif (item_scoreRF[i] == '2' or item_scoreRF[i] == '3' or item_scoreRF[i] == '4'):
				lista_option.append(item_scoreRF[i])
				i+=1
			else:
				lista_option.append(None)
				i+=1
		for q in queryset:
			if (lista_option[j] == True or lista_option[j] == False):		
				q.option = lista_option[j]
				j+=1
			else:
				q.extra_option = lista_option[j]
				j+=1

	return queryset



def set_option_to_item(mchat_item,item_score):
	i = 0
	if (len(item_score) and len(item_score) == mchat_item.count()):
		for m in mchat_item:
			if (item_score[i] == '1'):
				m.option = True
				i+=1
			else:
				m.option = False
				i+=1
		return mchat_item
	else:
		return mchat_item
	return mchat_item

def count_positive_mchat():
	dict_n = {
	  "n_patient_p": 0,
	  "n_patient_pc": 0,
	  "n_mchat_done": 0,
	  "ratio": 0
	  
	}
	n_patients_positive = Patient.objects.filter(positive=True).count()
	n_patients_positive_confirmed = Patient.objects.filter(positive_tr=True).count()
	n_mchat_done = Patient.objects.filter(finish=True).count()
	if n_patients_positive > 0:

		ratio = n_patients_positive_confirmed / n_patients_positive
	else:
		ratio = 0
	

	dict_n["n_patient_p"] = n_patients_positive
	dict_n["n_patient_pc"] = n_patients_positive_confirmed
	dict_n["n_mchat_done"] = n_mchat_done
	dict_n["ratio"] = ratio
	return dict_n

def audit_info_mapper(audit_int):
	message=""
	AUDIT_INFO = {
    2: 'Audicion normal',
    3: 'Audicion por debajo de lo normal',
    4: 'Resultados no concluyentes'
	}
	if audit_int == 'No aplica' or audit_int == '':
		message = "No realizado"
		return message
	else:
		message=AUDIT_INFO[int(audit_int)]
		return message

	return message


def generate_list_dict(Item_list):
	list_dict = []

	for i in Item_list:
		list_dict.append({"question_id":str(i.question_id),"static":"testM/item"+str(i.question_id)+".png"})
	return list_dict


	


@login_required
def mchat_start (request, pk):
	total_score=0
	listaFollowUp = ""
	lista_score = ""
	contador=0
	request.session['form_error'] = 0
	
	
	item = Item.objects.order_by('question_id') # Ordeno en base al numero de item	
	patient=Patient.objects.get(pk=pk)

	mchatFormSet = formset_factory(mchatTest ,extra=0)
	if request.method == "POST":
		print("entro en post")
		formset = mchatFormSet(request.POST,initial=[{'question': l.question,'question_id': l.question_id} for l in item])#debo pasar initial para que se sigan mostrando los datos en el caso de que no sea valido
		if formset.is_valid():
			print("valido")							
			for f in formset:
				option=f.cleaned_data['option']
				lista_score=option_to_score_list(option,lista_score)
				question_id=int(f.cleaned_data['question_id'])
				total_score+=MchatScore(question_id,option)				
				if(MchatScore(question_id,option) == 1):
					listaFollowUp=(construct_followup_list(question_id,listaFollowUp))
			dictr = IfMchatFollowUp(total_score)					
			Patient.objects.update_or_create(pk = pk, defaults={'mchat_score': total_score,'item_score': lista_score,'followup_list': listaFollowUp,'finish': dictr["finish"],'positive': dictr["result"],'item_scoreRF' : '','mchatrf_score': 0,'custom_quest': '','adic_info': ''})
			if(dictr["finish"] == True):
				Patient_historic.objects.create(patient = patient, mchat_score = total_score, item_score = lista_score, followup_list = listaFollowUp, positive = dictr["result"] )

			#return render(request,'testM/resultados_mchat.html', {'formset': formset,'total_score': total_score})
			print(dictr["result"])
			return render(request,'testM/resultados_mchat.html',{'total_score': total_score, 'dictr': dictr, 'pk': pk, 'patient': patient})
		else:
			request.session['form_error'] = formset.total_error_count()
			print("no es valido pinto esto")
			print(request.session['form_error'])
			print(formset.errors)

	else:
		print("estoy en el get")			
		formset = mchatFormSet(initial=[{'question': l.question,'question_id': l.question_id} for l in item])
	return render(request,'testM/mchatStart.html', {'formset': formset, 'patient': patient})


@login_required
def followup_mchat(request, pk):
	patient=Patient.objects.get(pk=pk)
	item_score=patient.item_score
	total_score = patient.mchat_score
	followup_array=patient.followup_list
	c_quest = patient.custom_quest
	followup_object=[]
	objects = ""
	list_pk_followup = []
	count_pasa = 0
	count_nopasa = 0
	count_group = 0
	list_trigger = []
	list_single = []
	item_scoreRF = ""

	

	#transformo el char followup_array a lista
	followup_list=char_to_list(followup_array)

	#transformo la cuestiones personalizadas a lista
	c_quest = c_quest.split(', ')

	#recorro la lista anterior y genero la lista de item que tienen seguimiento
	followup_object=objetc_to_list(followup_list)

	#utilizo la lista de item para filtrar los objetos followup que corresponden
	followUpItem=FollowUpItem.objects.filter(item__in=followup_object)

	#genero queryset item para paginacion 
	item=Item.objects.order_by('pk').filter(question_id__in=followup_list)
	
	mchatFollowUpFormset = formset_factory(mchatFollowup ,extra=0)
	customform = custom_question()
	#print(request.session.items())
	


	if 'validate' in request.POST:
		data_des = request.session['page_query']
		for obj in serializers.deserialize("json", data_des):
			list_pk_followup.append(obj.object.pk)
		followUpItem=FollowUpItem.objects.filter(pk__in=list_pk_followup)
		formset = mchatFollowUpFormset(request.POST,initial=[{'question_group': l.question_group,'question_item':l.question_item, 'question': l.question, 'extra_option': l.extra_option} for l in followUpItem])
		if formset.is_valid():
			for f in formset:
				option=f.cleaned_data['option']
				group=f.cleaned_data['question_group']
				item_scoreRF=option_to_score_list(option,item_scoreRF)
				request.session['item_scoreRF'] = item_scoreRF
				if (group == TRIGGER):
					list_trigger.append(option_to_n(option))					

				elif (group == SINGLE):
					list_single.append(option_to_n(option))

				elif (group == GROUP):
					count_group+=option_to_n(option)

				elif (group == PASA):
					count_pasa+=option_to_n(option)

				elif (group == NOPASA):
					count_nopasa+=option_to_n(option)
				elif (group == AUDIT):
					extra_option = f.cleaned_data['extra_option']
					top=len(request.session['item_scoreRF'])
					request.session['item_scoreRF']= request.session['item_scoreRF'][(0):(top-1)]
					if (extra_option != ''):
						request.session['item_scoreRF']+= extra_option
					else:
						request.session['item_scoreRF']+= 'X'
			question_id = request.session['question_id']			
			request.session['score_rf']+= MchatScoreRf(question_id,count_pasa,list_trigger,count_nopasa,count_group,list_single)
			request.session['string_score'] += str(MchatScoreRf(question_id,count_pasa,list_trigger,count_nopasa,count_group,list_single))
			if (question_id == 2):
				request.session['audit_info'] = extra_option

			n_pages=request.session['num_pages']
			a_page=request.session['actual_page']
			request.session[str(a_page)] = request.session['item_scoreRF']
			request.session['post_send'] = a_page + 1
			if(a_page + 1 <= n_pages):
				return redirect('/mchat/followup/' + str(pk) + '?page=' + str(a_page+1))
			else:
				score_rf = request.session['score_rf']
				audit_info = request.session['audit_info']
				positive=FinalMchatRfScore(score_rf)
				item_risk = request.session['string_score']
				item_scoreRF = add_all_optionRF(n_pages,request)
				custom_quest = patient.custom_quest
				Patient.objects.update_or_create(pk = pk, defaults={'positive': positive,'audit_info': audit_info,'item_scoreRF': item_scoreRF,'mchatrf_score': score_rf,'followup_risk':item_risk,'finish': True,'adic_info' : ''})
				Patient_historic.objects.create(patient = patient, mchat_score = total_score, 
					mchatrf_score = score_rf , item_score = item_score, item_scoreRF = item_scoreRF, 
					followup_list = followup_array, followup_risk = item_risk, positive = positive, 
					audit_info = audit_info, custom_quest = custom_quest, adic_info = '')
				return render(request,'testM/resultados_mchatRF.html',{'score_rf': score_rf, 'positive': positive, 'patient': patient,})
			return redirect('mchats:mchats')
		else:
			print("Not Valid")
			print(formset.errors)

	else:
		print("esto es get")
		paginator = Paginator(item, 1)
		page = request.GET.get('page')
		if(page == None):
			request.session['post_send'] = 1
			request.session['actual_page']= 1
			request.session['score_rf'] = 0
			request.session['string_score'] = ""
			request.session['audit_info'] = "No aplica"
			request.session['item_scoreRF'] = ""
			request.session['num_elem_rf'] =""
		else:
			request.session['actual_page']= int(page)
		request.session['num_pages']=paginator.num_pages		

		try:
			objects = paginator.page(page)

		except PageNotAnInteger:
			objects = paginator.page(1)

		except EmptyPage:
			objects = paginator.page(paginator.num_pages)
		#creo otro filtro para los item de siguimiento y asi en cada pagina solo salen los del item correspondiente
		page_query = followUpItem.filter(item__pk__in=[object.pk for object in objects])
		print(request.session['num_elem_rf'])
		print("actual page:" + str(request.session['actual_page']))
		print("siguiente:" + str(request.session['post_send']))
		if request.session.get(str(page)) is not None:
			print(request.session[str(page)])
		else:
			request.session[str(page)] = ""

		
		if (request.session['actual_page'] < request.session['post_send']):
			print(request.session[str(page)])
			if ((len(request.session['string_score']) > 0) and (len(request.session['string_score'])) == (request.session['actual_page'])):
				if(request.session['string_score'][request.session['actual_page'] - 1] == '1'):
					request.session['score_rf']-=1
				request.session['string_score']=request.session['string_score'][0:request.session['actual_page']-1]
				print("socre")
		print("String Score: "+request.session['string_score'])
		print(request.session['score_rf'])

		for o in objects:
			request.session['question_id'] = o.question_id

		page_query=set_option_to_rf_nf(page_query,request.session[str(page)])
		element=select_c_quest(c_quest,request.session['question_id'])
		print(element)

		formset = mchatFollowUpFormset(initial=[{'question_group': l.question_group,'question_item':l.question_item, 'question': l.question,'option': l.option, 'extra_option': l.extra_option} for l in page_query])
		data = serializers.serialize("json", page_query)
		request.session['page_query'] = data

	return render(request,'testM/followUp_mchat.html',{'patient': patient,'formset': formset,'objects': objects,'customform': customform,'element':element,})

@login_required
def save_custom_quest(request,pk,question_id):
	patient = get_object_or_404(Patient,pk=pk)
	question = patient.custom_quest
	message = ''
	customform = custom_question(request.POST)
	if customform.is_valid():
		question_form = customform.cleaned_data.get('custom_quest')
		if(split_validate(question,question_form,question_id) == True):
			question = parse_custom_quest(question,question_form,question_id)
			Patient.objects.update_or_create(pk = pk, defaults={'custom_quest': question,})
			message = "Se añadió: " + question[2:(len(question))] + " como pregunta personalizada"
		else:
			message = 'Ya introdujo esta pregunta anteriormente'
	return JsonResponse(

			{
				'content' : {

						'message' : message,
				}
			}
		)

@login_required
def undo_custom_quest(request,pk,question_id):
	patient = get_object_or_404(Patient,pk=pk)
	question = patient.custom_quest
	message = "No se envió ninguna pregunta personalizada"
	if len(question) > 0:
		question = question.split(', ')
		if split_validate_undo(question,question_id) == True:
			new_list = undo_custom(question,question_id)
			Patient.objects.update_or_create(pk = pk, defaults={'custom_quest': new_list,})
			message = 'Se ha descartado correctamente'
		else:
			message = 'Ninguna pregunta que descartar'
	return JsonResponse(

			{
				'content' : {

						'message' : message,
				}
			}
		)



@login_required
def patient_result(request,pk):
	patient = Patient.objects.get(pk=pk)
	item_score = patient.item_score
	item_scoreRF = patient.item_scoreRF
	followup_array = patient.followup_list
	question_c = patient.custom_quest
	item_risk = patient.followup_risk
	adic_info = patient.adic_info
	mchat_item = Item.objects.all()
	audit_info = patient.audit_info
	last_test = Patient_historic.objects.filter(patient = pk).last()
	delta = date.today() - patient.birth_date
	delta = delta.days
	age = transform_date_to_age(delta)
	if last_test != None:

		last_test = last_test.date_test
	
	audit_message = ""
	Item_list = []
	target=""
	namepdf=""

	#transformo el char followup_array a lista
	followup_list = char_to_list(followup_array)
	
	#recorro la lista anterior y genero la lista de item que tienen seguimiento
	Item_list = objetc_to_list(followup_list)

	#utilizo la lista de item para filtrar los objetos followup que corresponden
	followUpItem = FollowUpItem.objects.filter(item__in=Item_list)
	#Transformo el campo audit_info que corresponde con un tipo númerico a su correspondiente mensaje del nivel auditivo
	audit_message=audit_info_mapper(audit_info)

	followUpItem = set_option_to_rf(followUpItem,question_c,item_scoreRF)

	mchat_item = set_option_to_item(mchat_item,item_score)

	Item_dict = generate_list_dict(Item_list)

	List_risk = item_risk2list(item_risk,Item_list)

	form = adicional_info()

	adic_info_list = adic_info.split("||")


	if request.method == 'POST':
		html_string = render_to_string('testM/patient_result.html', {'followUpItem': followUpItem,'mchat_item': mchat_item,'patient': patient,'audit_message': audit_message,'last_test': last_test,'age': age,'Item_dict': Item_dict, 'List_risk' : List_risk, 'adic_info_list' : adic_info_list},request=request)
		html = HTML(string=html_string,base_url=request.build_absolute_uri())
		namepdf = "mchat_" + str(patient).replace(" ","_") + ".pdf"
		target = "/tmp/" + namepdf
		css_file = '/testM/static/css/w3.css'
		html.write_pdf(target=target,stylesheets=[CSS(settings.BASE_DIR +  css_file)],presentational_hints=True)

		fs = FileSystemStorage('/tmp')
		with fs.open(namepdf) as pdf:
			response = HttpResponse(pdf, content_type='application/pdf')
			response['Content-Disposition'] = 'attachment; filename=' + namepdf
			time.sleep(1)
			return response
		return response

	return render(request, 'testM/patient_result.html',{'followUpItem': followUpItem,'mchat_item': mchat_item,
		'patient': patient,'audit_message': audit_message,'last_test': last_test,
		'age': age, 'Item_dict': Item_dict,'List_risk' : List_risk, 'form' : form, 'adic_info_list' : adic_info_list})


@login_required
def patient_historic_result(request,pk):
	patient = Patient_historic.objects.get(pk=pk)
	item_score = patient.item_score
	item_scoreRF = patient.item_scoreRF
	item_risk = patient.followup_risk
	followup_array = patient.followup_list
	question_c = patient.custom_quest
	birth_date = patient.patient.birth_date
	adic_info = patient.adic_info
	date_test = patient.date_test
	mchat_item = Item.objects.all()
	audit_info = patient.audit_info
	audit_message = ""
	Item_list = []
	target = ""
	namepdf = ""
	#Calculo la edad que tenia el dia de la realización del test
	delta = date_test.date() - birth_date
	delta = delta.days
	age = transform_date_to_age(delta)

	#transformo el char followup_array a lista
	followup_list = char_to_list(followup_array)
	
	#recorro la lista anterior y genero la lista de item que tienen seguimiento
	Item_list = objetc_to_list(followup_list)

	#utilizo la lista de item para filtrar los objetos followup que corresponden
	followUpItem = FollowUpItem.objects.filter(item__in=Item_list)
	#Transformo el campo audit_info que corresponde con un tipo númerico a su correspondiente mensaje del nivel auditivo
	audit_message=audit_info_mapper(audit_info)

	followUpItem = set_option_to_rf(followUpItem,question_c,item_scoreRF)

	mchat_item = set_option_to_item(mchat_item,item_score)

	Item_dict = generate_list_dict(Item_list)

	List_risk = item_risk2list(item_risk,Item_list)

	form = adicional_info()

	adic_info_list = adic_info.split("||")


	if request.method == 'POST':
		html_string = render_to_string('testM/patient_historic_result.html', {'followUpItem': followUpItem,'mchat_item': mchat_item,'patient': patient,'audit_message': audit_message,'age':age,'Item_dict':Item_dict, 'List_risk' : List_risk , 'adic_info_list':adic_info_list},request=request)
		html = HTML(string=html_string,base_url=request.build_absolute_uri())
		namepdf = "mchat_" + str(patient.patient).replace(" ","_") + str(patient.date_test) + ".pdf"
		target = "/tmp/" + namepdf
		css_file = '/testM/static/css/w3.css'
		html.write_pdf(target=target,stylesheets=[CSS(settings.BASE_DIR +  css_file)],presentational_hints=True);

		fs = FileSystemStorage('/tmp')
		with fs.open(namepdf) as pdf:
			response = HttpResponse(pdf, content_type='application/pdf')
			response['Content-Disposition'] = 'attachment; filename=' + namepdf
			return response
		return response


	return render(request, 'testM/patient_historic_result.html',{'followUpItem': followUpItem,'mchat_item': mchat_item,
		'patient': patient,'audit_message': audit_message,'age': age,
		'Item_dict':Item_dict,'List_risk' : List_risk, 'form' : form, 'adic_info_list': adic_info_list})
@login_required
def graphics(request):
	dict_n = count_positive_mchat()

	return render(request, 'testM/graphic_mchat.html',{'dict_n': dict_n})

@login_required
def confirm_positive(request,pk):
	positive_tr = Patient.objects.get(pk=pk).positive_tr
	if positive_tr == True:
		Patient.objects.update_or_create(pk = pk, defaults={'positive_tr': False,})
	else:
		Patient.objects.update_or_create(pk = pk, defaults={'positive_tr': True,})
	return redirect('mchats:patients')


def guide_mchat(request):
	return render(request, 'testM/guide_mchat.html')


def result_docx(request,pk):
	patient = Patient.objects.get(pk=pk)
	item_score = patient.item_score
	item_scoreRF = patient.item_scoreRF
	followup_array = patient.followup_list
	question_c = patient.custom_quest
	item_risk = patient.followup_risk
	mchat_item = Item.objects.all()
	audit_info = patient.audit_info
	adic_info = patient.adic_info
	last_test = Patient_historic.objects.filter(patient = pk).last()
	delta = date.today() - patient.birth_date
	delta = delta.days
	age = transform_date_to_age(delta)
	if last_test != None:
		last_test = last_test.date_test
	
	audit_message = ""
	Item_list = []
	target=""
	namepdf=""

	#transformo el char followup_array a lista
	followup_list = char_to_list(followup_array)
	
	#recorro la lista anterior y genero la lista de item que tienen seguimiento
	Item_list = objetc_to_list(followup_list)

	#utilizo la lista de item para filtrar los objetos followup que corresponden
	followUpItem = FollowUpItem.objects.filter(item__in=Item_list)
	#Transformo el campo audit_info que corresponde con un tipo númerico a su correspondiente mensaje del nivel auditivo
	audit_message=audit_info_mapper(audit_info)

	followUpItem = set_option_to_rf(followUpItem,question_c,item_scoreRF)

	mchat_item = set_option_to_item(mchat_item,item_score)

	Item_dict = generate_list_dict(Item_list)

	List_risk = item_risk2list(item_risk,Item_list)

	adic_info_list = adic_info.split("||")

	document = Document()

	namedocx = "mchat_" + str(patient).replace(" ","_") + ".docx"
	docx_title = namedocx

	document.add_heading("Resultados M-CHAT-R/F",0)

	document.add_heading("Resultados Obtenidos",level=1)

	document.add_paragraph()



	a = document.add_paragraph()
	b = document.add_paragraph()
	c = document.add_paragraph()
	d = document.add_paragraph()
	e = document.add_paragraph()
	f = document.add_paragraph()
	g = document.add_paragraph()

	a.add_run(str(patient)).bold = True
	b.add_run('Puntuación M-CHAT-R: ').bold = True
	b.add_run(str(patient.mchat_score))
	c.add_run('Puntuación M-CHAT-R/F: ').bold = True
	c.add_run(str(patient.mchatrf_score))
	d.add_run('Nivel de audición: ').bold = True
	d.add_run(audit_message)
	e.add_run('Positivo en el test M-CHAT-R/F: ').bold = True
	if (patient.positive):
		e.add_run('Si')
	else:
		e.add_run('No')	
	f.add_run('Fecha de realización del último test: ').bold = True
	f.add_run(str(last_test))
	g.add_run('Edad Actual: ').bold = True
	g.add_run(age)


	document.add_heading("Cuestiones M-CHAT-R",level=1)
	document.add_heading()

	for m in mchat_item:
		if m.option == True:
			document.add_paragraph(m.question + ' - Usted respondió: Si ',style='List Bullet')
		else:
			document.add_paragraph(m.question + ' - Usted respondió: No ',style='List Bullet')

	document.add_heading("Cuestiones M-CHAT-R/F",level=1)
	document.add_heading()

	if patient.mchat_score > 2 and len(patient.item_scoreRF) > 0:
		for f in followUpItem:			
			if f.option == True and f.question_group != 'Audit':
				document.add_paragraph(f.question + ' - Usted respondió: Si / Item: ' + str(f.item),style='List Bullet')			
			elif f.option == False and f.question != 'Audit':
				document.add_paragraph(f.question + ' - Usted respondió: No / Item: ' + str(f.item),style='List Bullet')
			else:
				document.add_paragraph(f.question + ' - Usted respondió:' + audit_message + '/ Item: ' + str(f.item),style='List Bullet')
		document.add_heading("Flujos que implican riesgo de TEA",level=1)
		document.add_paragraph()

		if (len(List_risk) > 0):
			for l in List_risk:
				document.add_paragraph("Item: "+str(l.question_id),style='List Bullet')
		else:
			document.add_paragraph("Ningún item implica riesgo de TEA")

		document.add_heading("Flujos realizados",level=1)
		document.add_paragraph()
		for t in Item_dict:
			document.add_paragraph("Flujo del Item "+t["question_id"])
			document.add_picture(settings.BASE_DIR+"/testM/static/testM/item"+t["question_id"]+".png", width=Inches(4))	
	elif( patient.mchat_score > 2 and patient.mchat_score < 8 and len(patient.item_scoreRF) == 0):
		document.add_paragraph('No ha realizado aún la entrevista de seguimiento')
	else:
		document.add_paragraph()
		document.add_paragraph('No fue necesario relizar la entrevista de seguimiento')
	document.add_heading("Contenido Adicional", level=1)
	if adic_info != '':
		document.add_paragraph()
		for l in adic_info_list:
			document.add_paragraph(l, style='List Bullet')
	else:
		document.add_paragraph("No se ha añadido contenido adicional", style='List Bullet')



	return generate_docx(docx_title,document)


def result_docx_historic(request,pk):
	patient = Patient_historic.objects.get(pk=pk)
	item_score = patient.item_score
	item_scoreRF = patient.item_scoreRF
	item_risk = patient.followup_risk
	followup_array = patient.followup_list
	question_c = patient.custom_quest
	birth_date = patient.patient.birth_date
	date_test = patient.date_test
	mchat_item = Item.objects.all()
	audit_info = patient.audit_info
	adic_info = patient.adic_info
	audit_message = ""
	Item_list = []
	target = ""
	namepdf = ""
	#Calculo la edad que tenia el dia de la realización del test
	delta = date_test.date() - birth_date
	delta = delta.days
	age = transform_date_to_age(delta)

	#transformo el char followup_array a lista
	followup_list = char_to_list(followup_array)
	
	#recorro la lista anterior y genero la lista de item que tienen seguimiento
	Item_list = objetc_to_list(followup_list)

	#utilizo la lista de item para filtrar los objetos followup que corresponden
	followUpItem = FollowUpItem.objects.filter(item__in=Item_list)
	#Transformo el campo audit_info que corresponde con un tipo númerico a su correspondiente mensaje del nivel auditivo
	audit_message=audit_info_mapper(audit_info)

	followUpItem = set_option_to_rf(followUpItem,question_c,item_scoreRF)

	mchat_item = set_option_to_item(mchat_item,item_score)

	Item_dict = generate_list_dict(Item_list)

	List_risk = item_risk2list(item_risk,Item_list)

	adic_info_list = adic_info.split("||")

	document = Document()

	namedocx = "mchat_" + str(patient).replace(" ","_") + ".docx"
	docx_title = namedocx

	document.add_heading("Resultados M-CHAT-R/F",0)

	document.add_heading("Resultados Obtenidos",level=1)

	document.add_paragraph()



	a = document.add_paragraph()
	b = document.add_paragraph()
	c = document.add_paragraph()
	d = document.add_paragraph()
	e = document.add_paragraph()
	f = document.add_paragraph()
	g = document.add_paragraph()

	a.add_run(str(patient.patient)).bold = True
	b.add_run('Puntuación M-CHAT-R: ').bold = True
	b.add_run(str(patient.mchat_score))
	c.add_run('Puntuación M-CHAT-R/F: ').bold = True
	c.add_run(str(patient.mchatrf_score))
	d.add_run('Nivel de audición: ').bold = True
	d.add_run(audit_message)
	e.add_run('Positivo en el test M-CHAT-R/F: ').bold = True
	if (patient.positive):
		e.add_run('Si')
	else:
		e.add_run('No')	
	f.add_run('Fecha de realización: ').bold = True
	f.add_run(str(date_test))
	g.add_run('Edad el día de la realización: ').bold = True
	g.add_run(age)


	document.add_heading("Cuestiones M-CHAT-R",level=1)
	document.add_heading()

	for m in mchat_item:
		if m.option == True:
			document.add_paragraph(m.question + ' - Usted respondió: Si ',style='List Bullet')
		else:
			document.add_paragraph(m.question + ' - Usted respondió: No ',style='List Bullet')

	document.add_heading("Cuestiones M-CHAT-R/F",level=1)
	document.add_heading()

	if patient.mchat_score > 2 and len(patient.item_scoreRF) > 0:
		for f in followUpItem:			
			if f.option == True and f.question_group != 'Audit':
				document.add_paragraph(f.question + ' - Usted respondió: Si / Item: ' + str(f.item),style='List Bullet')			
			elif f.option == False and f.question != 'Audit':
				document.add_paragraph(f.question + ' - Usted respondió: No / Item: ' + str(f.item),style='List Bullet')
			else:
				document.add_paragraph(f.question + ' - Usted respondió:' + audit_message + '/ Item: ' + str(f.item),style='List Bullet')

		document.add_heading("Flujos que implican riesgo de TEA",level=1)
		document.add_paragraph()

		if (len(List_risk) > 0):
			for l in List_risk:
				document.add_paragraph("Item: "+str(l.question_id),style='List Bullet')
		else:
			document.add_paragraph("Ningún item implica riesgo de TEA")

		document.add_heading("Flujos realizados",level=1)
		document.add_paragraph()
		for t in Item_dict:
			document.add_paragraph("Flujo del Item "+t["question_id"])
			document.add_picture(settings.BASE_DIR+"/testM/static/testM/item"+t["question_id"]+".png", width=Inches(4))
	elif( patient.mchat_score > 2 and patient.mchat_score < 8 and len(patient.item_scoreRF) == 0):
		document.add_paragraph('No ha realizado aún la entrevista de seguimiento')
	else:
		document.add_paragraph()
		document.add_paragraph('No fue necesario relizar la entrevista de seguimiento')
	document.add_heading("Contenido Adicional", level=1)
	if adic_info != '':
		document.add_paragraph()
		for l in adic_info_list:
			document.add_paragraph(l, style='List Bullet')
	else:
		document.add_paragraph("No se ha añadido contenido adicional", style='List Bullet')


	return generate_docx(docx_title,document)


def generate_docx(docx,document):
	f = BytesIO()
	document.save(f)
	length = f.tell()
	f.seek(0)
	response = HttpResponse(
		f.getvalue(),
		content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'

	)
	response['Content-Disposition'] = 'attachment; filename=' + docx
	response['Content-Length'] = length
	return response

def save_adicional_info(request,pk):
	if request.META["HTTP_REFERER"].find("patientHistoric") == -1:
		patient = get_object_or_404(Patient,pk=pk)
	else:
		patient = get_object_or_404(Patient_historic, pk=pk)
	adic_info = patient.adic_info
	message = "Guardado correctamente"

	form = adicional_info(request.POST)
	if form.is_valid():
		info = form.cleaned_data.get('adic_info')
		if adic_info != '':
			adic_info+= '||' + info
		else:
			adic_info = info + '. '
		if request.META["HTTP_REFERER"].find("patientHistoric") == -1:
			Patient.objects.update_or_create(pk = pk, defaults={'adic_info': adic_info,})
		else:
			Patient_historic.objects.update_or_create(pk = pk, defaults={'adic_info': adic_info,})		
		return JsonResponse(

			{
				'content' : {

						'message' : message + " Recargue la página para ver los cambios",
				}
			}
		)
	else:
		print(form.errors.as_data()['adic_info'])
		message = str(form.errors.as_data()['adic_info'][0])
	return JsonResponse(

			{
				'content' : {

						'message' : message,
				}
			}
		)





def activation_sent_view(request):
	return render(request, 'activation_sent.html')


def activate(request, uidb64, token):
	try:
		uid = force_text(urlsafe_base64_decode(uidb64))
		user = User.objects.get(pk=uid)
		profile = Profile.objects.get(user__in = [user.pk])
	except (TypeError, ValueError, OverflowError, User.DoesNotExist):
		user = None

	if user is not None and account_activation_token.check_token(user,token):

		user.is_active = True

		user.profile.signup_confirmation = True
		user.save()
		login(request, user)
		url=reverse('mchats:profile_update',kwargs={'pk':user.pk})
		return HttpResponseRedirect(url)
	else:
		return render(request, 'activation_invalid.html')



def signup(request):
	if request.method == 'POST':
		form = SignUpForm(request.POST)
		if form.is_valid():
			user = form.save()
			user.refresh_from_db()
			user.profile.first_name = form.cleaned_data.get('first_name')
			user.profile.last_name = form.cleaned_data.get('last_name')
			user.profile.email = form.cleaned_data.get('email')
			user.profile.birth_date = form.cleaned_data.get('birth_date')
			#No podemos loguear al usuario hasta que confirme
			user.is_active = False
			user.save()
			current_site = get_current_site(request)
			subject  = 'M-CHAT HELP WEB APP: Por favor es necesario activar tu cuenta'
			email_from = settings.EMAIL_HOST_USER
			recipient_list = [user.profile.email,]
			message = render_to_string('activation_request.html', {
				'user': user,
				'domain': current_site.domain,
				'uid': urlsafe_base64_encode(force_bytes(user.pk)),
				'token': account_activation_token.make_token(user),
			})
			
			send_mail(subject, None, email_from, recipient_list,html_message=message)
			return redirect('mchats:activation_sent')
	else:
		form = SignUpForm()
	return render(request, 'signup.html',{'form':form})